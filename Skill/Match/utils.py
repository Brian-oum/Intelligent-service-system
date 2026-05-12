from .models import ServiceProvider
import math

def find_best_company(service_request):
    """
    Matches a user's service request to the best available verified company.
    Based on service category, location, and rating.
    """
    # Step 1: Filter by service category
    companies = ServiceProvider.objects.filter(
        service_category__icontains=service_request.service_type,
        verified=True
    )

    # Step 2: Filter by location match
    companies = companies.filter(location__icontains=service_request.location)

    # Step 3: Sort by rating (highest first)
    companies = companies.order_by('-rating')

    # Step 4: Pick best match
    return companies.first() if companies.exists() else None


def haversine_distance(lat1, lon1, lat2, lon2):
    """
    Calculate distance between two coordinates in KM
    """

    R = 6371  # Earth radius in KM

    lat1 = math.radians(float(lat1))
    lon1 = math.radians(float(lon1))
    lat2 = math.radians(float(lat2))
    lon2 = math.radians(float(lon2))

    dlon = lon2 - lon1
    dlat = lat2 - lat1

    a = math.sin(dlat / 2)**2 + math.cos(lat1) * math.cos(lat2) * math.sin(dlon / 2)**2
    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))

    return R * c

# utils.py
from django.core.mail import send_mail
from django.conf import settings

def send_notification_email(subject, message, recipient_email):
    send_mail(
        subject,
        message,
        settings.DEFAULT_FROM_EMAIL,  # e.g., "noreply@yourdomain.com"
        [recipient_email],
        fail_silently=False
    )

"""
reports/utils.py
────────────────
Generates the Service Provider .docx report from the context dict
produced by build_report_context() in views.py.

Requires:  pip install docx
"""

import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from django.utils.timezone import localtime


# ─── Colour constants (hex strings, no #) ────────────────────────────────────
BLUE_DARK  = "1A3C6E"
BLUE_MID   = "2E6DB4"
BLUE_LIGHT = "D6E4F0"
ACCENT     = "27AE60"
GREY_BG    = "F4F6F9"
GREY_TEXT  = "555555"
WHITE      = "FFFFFF"


def _hex_to_rgb(hex_str):
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ─── Low-level XML helpers ────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color.lstrip('#'))
    tcPr.append(shd)


def _set_cell_borders(cell, color="CCCCCC", size="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    size)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color.lstrip('#'))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _remove_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_table_width(table, width_dxa):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(width_dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def _set_col_width(cell, width_dxa):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _add_run(para, text, bold=False, size=10, color=None, italic=False):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = _hex_to_rgb(color)
    return run


# ─── High-level helpers ───────────────────────────────────────────────────────

def _add_section_header(doc, title):
    """Full-width dark-blue section heading row."""
    table = doc.add_table(rows=1, cols=1)
    _set_table_width(table, 9360)
    cell = table.cell(0, 0)
    _set_col_width(cell, 9360)
    _set_cell_bg(cell, BLUE_DARK)
    _remove_cell_borders(cell)
    _set_cell_margins(cell, top=120, bottom=120, left=200, right=200)

    para = cell.paragraphs[0]
    _add_run(para, title, bold=True, size=12, color=WHITE)
    doc.add_paragraph()   # spacer


def _add_kv_table(doc, rows):
    """Two-column label → value table."""
    table = doc.add_table(rows=len(rows), cols=2)
    _set_table_width(table, 9360)

    for i, (label, value) in enumerate(rows):
        bg = GREY_BG if i % 2 == 0 else WHITE

        lc = table.cell(i, 0)
        _set_col_width(lc, 2800)
        _set_cell_bg(lc, bg)
        _set_cell_borders(lc)
        _set_cell_margins(lc)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = lc.paragraphs[0]
        _add_run(lp, str(label), bold=True, size=10, color=BLUE_DARK)

        vc = table.cell(i, 1)
        _set_col_width(vc, 6560)
        _set_cell_bg(vc, bg)
        _set_cell_borders(vc)
        _set_cell_margins(vc)
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vp = vc.paragraphs[0]
        _add_run(vp, str(value), size=10, color=GREY_TEXT)

    doc.add_paragraph()


def _add_data_table(doc, headers, rows, col_widths):
    """Generic multi-column table with a blue header row."""
    total_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=total_cols)
    _set_table_width(table, sum(col_widths))

    # Header row
    for ci, (h, w) in enumerate(zip(headers, col_widths)):
        cell = table.cell(0, ci)
        _set_col_width(cell, w)
        _set_cell_bg(cell, BLUE_MID)
        _set_cell_borders(cell, color=BLUE_MID)
        _set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, h, bold=True, size=10, color=WHITE)

    # Data rows
    for ri, row in enumerate(rows):
        bg = WHITE if ri % 2 == 0 else GREY_BG
        for ci, (val, w) in enumerate(zip(row, col_widths)):
            cell = table.cell(ri + 1, ci)
            _set_col_width(cell, w)
            _set_cell_bg(cell, bg)
            _set_cell_borders(cell)
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p, str(val), size=9.5, color=GREY_TEXT)

    doc.add_paragraph()


def _add_metric_row(doc, metrics):
    """Single row of KPI tiles."""
    n = len(metrics)
    col_w = 9360 // n
    table = doc.add_table(rows=1, cols=n)
    _set_table_width(table, col_w * n)

    for i, m in enumerate(metrics):
        cell = table.cell(0, i)
        _set_col_width(cell, col_w)
        _set_cell_bg(cell, BLUE_LIGHT)
        _set_cell_borders(cell, color=BLUE_MID)
        _set_cell_margins(cell, top=160, bottom=160, left=160, right=160)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        vp = cell.add_paragraph()
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(vp, str(m['value']), bold=True, size=18, color=BLUE_DARK)

        lp = cell.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(lp, m['label'], size=9, color=GREY_TEXT)

        # remove the auto-created blank first paragraph
        cell._tc.remove(cell._tc.p_lst[0])

    doc.add_paragraph()


# ─── Stars helper ────────────────────────────────────────────────────────────

def _stars(rating):
    try:
        r = int(round(float(rating)))
        return ('★' * r) + ('☆' * (5 - r)) + f'  {rating}'
    except (ValueError, TypeError):
        return str(rating)


# ─── Main generator ──────────────────────────────────────────────────────────

def generate_provider_docx(context):
    """
    Accepts the context dict from build_report_context() and returns
    a bytes buffer containing the complete .docx file.
    """
    provider        = context['provider']
    report_date     = localtime(context['report_date']).strftime('%d %B %Y')
    services        = context['services']
    service_stats   = context['service_stats']
    recent_requests = context['recent_requests']
    recent_reviews  = context['recent_reviews']
    documents       = context['documents']
    status_breakdown = context['status_breakdown']

    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = section.right_margin  = Inches(0.75)
    section.top_margin    = section.bottom_margin = Inches(0.75)

    # ── Default font ──────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # ═══════════════════════════════════════════════════════════════════════════
    # COVER BANNER
    # ═══════════════════════════════════════════════════════════════════════════
    banner = doc.add_table(rows=1, cols=1)
    _set_table_width(banner, 9360)
    bc = banner.cell(0, 0)
    _set_col_width(bc, 9360)
    _set_cell_bg(bc, BLUE_DARK)
    _remove_cell_borders(bc)
    _set_cell_margins(bc, top=300, bottom=300, left=300, right=300)

    tp = bc.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(tp, provider.company_name, bold=True, size=22, color=WHITE)

    sp = bc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(sp, f"Service Provider Report  •  {report_date}", size=11, color="A8C4E2")

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # KPI TILES
    # ═══════════════════════════════════════════════════════════════════════════
    _add_metric_row(doc, [
        {'label': 'Total Services',    'value': context['total_services']},
        {'label': 'Total Requests',    'value': context['total_requests']},
        {'label': 'Completion Rate',   'value': f"{context['completion_rate']}%"},
        {'label': 'Average Rating',    'value': f"{context['avg_rating']} ★"},
    ])

    # ═══════════════════════════════════════════════════════════════════════════
    # 1. PROVIDER PROFILE
    # ═══════════════════════════════════════════════════════════════════════════
    _add_section_header(doc, "1.  Provider Profile")
    _add_kv_table(doc, [
        ("Company Name",      provider.company_name),
        ("Account Owner",     provider.user.get_full_name() or provider.user.username),
        ("Email",             provider.user.email),
        ("Contact Number",    provider.contact_number),
        ("Address",           provider.address),
        ("Website",           provider.website or "—"),
        ("Member Since",      provider.created_at.strftime('%d %B %Y') if provider.created_at else "—"),
        ("Profile Completed", "Yes" if provider.profile_completed else "No"),
        ("Admin Verified",    "✔ Verified" if provider.is_verified else "✘ Not Verified"),
        ("Account Status",    "Active" if provider.is_active else "Inactive"),
        ("GPS Coordinates",   f"{provider.latitude}, {provider.longitude}" if provider.latitude else "—"),
    ])

    # ═══════════════════════════════════════════════════════════════════════════
    # 2. SERVICES CATALOGUE
    # ═══════════════════════════════════════════════════════════════════════════
    _add_section_header(doc, "2.  Services Catalogue")
    _add_data_table(
        doc,
        headers   = ["Service Title", "Category", "Status", "Verified", "Requests", "Avg Rating"],
        rows      = [
            [
                s['service'].title,
                s['service'].category.name,
                "Active" if s['service'].is_active else "Inactive",
                "Yes" if s['service'].is_verified else "No",
                s['req_count'],
                _stars(s['avg_rating']),
            ]
            for s in service_stats
        ],
        col_widths = [2800, 1600, 1000, 900, 1100, 1960],
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # 3. RECENT SERVICE REQUESTS
    # ═══════════════════════════════════════════════════════════════════════════
    _add_section_header(doc, "3.  Recent Service Requests (Last 10)")
    _add_data_table(
        doc,
        headers   = ["#", "Service", "Client", "Location", "Date", "Status"],
        rows      = [
            [
                f"SR-{r.pk}",
                r.service.title,
                r.user.get_full_name() or r.user.username,
                r.location,
                r.created_at.strftime('%d %b %Y'),
                r.get_status_display(),
            ]
            for r in recent_requests
        ],
        col_widths = [700, 2100, 1600, 2000, 1160, 1200],
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # 4. CLIENT REVIEWS
    # ═══════════════════════════════════════════════════════════════════════════
    _add_section_header(doc, "4.  Client Reviews & Ratings (Latest 5)")
    _add_data_table(
        doc,
        headers   = ["Service", "Client", "Rating", "Comment"],
        rows      = [
            [
                rv.service_request.service.title,
                rv.user.get_full_name() or rv.user.username,
                _stars(rv.rating),
                rv.comment or "—",
            ]
            for rv in recent_reviews
        ],
        col_widths = [2400, 1700, 1500, 3760],
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # 5. VERIFICATION DOCUMENTS
    # ═══════════════════════════════════════════════════════════════════════════
    _add_section_header(doc, "5.  Verification Documents")
    _add_data_table(
        doc,
        headers   = ["Document Name", "Upload Date", "Status"],
        rows      = [
            [
                d.document_name,
                d.uploaded_at.strftime('%d %b %Y'),
                "Submitted ✔",
            ]
            for d in documents
        ] or [["No documents on file", "—", "—"]],
        col_widths = [5200, 2100, 2060],
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # 6. PERFORMANCE INSIGHTS
    # ═══════════════════════════════════════════════════════════════════════════
    _add_section_header(doc, "6.  Performance Insights")

    sub = doc.add_paragraph()
    _add_run(sub, "Request Status Breakdown", bold=True, size=11, color=BLUE_DARK)

    _add_data_table(
        doc,
        headers   = ["Status", "Count", "Percentage"],
        rows      = [
            [s['label'], s['count'], f"{s['pct']}%"]
            for s in status_breakdown
        ],
        col_widths = [3120, 3120, 3120],
    )

    sub2 = doc.add_paragraph()
    _add_run(sub2, "Top-Performing Services", bold=True, size=11, color=BLUE_DARK)

    _add_data_table(
        doc,
        headers   = ["Rank", "Service Title", "Total Requests", "Avg Rating"],
        rows      = [
            [i + 1, s['service'].title, s['req_count'], _stars(s['avg_rating'])]
            for i, s in enumerate(service_stats[:5])
        ],
        col_widths = [700, 4500, 2100, 2060],
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # FOOTER NOTE
    # ═══════════════════════════════════════════════════════════════════════════
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(
        footer_p,
        "This report is generated automatically by the HomeFix Connect platform "
        "and is intended solely for authorised administrators.",
        italic=True, size=8, color=GREY_TEXT,
    )
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(date_p, f"Report Date: {report_date}", size=8, color=GREY_TEXT)

    # ── Serialise to buffer ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()